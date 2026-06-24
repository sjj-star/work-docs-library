"""office_parser 模块."""

import hashlib
from pathlib import Path

from core.models import Chapter, Document


class OfficeParser:
    """OfficeParser 类.

    DOCX/XLSX 解析当前未接入主 pipeline，需要额外安装 office 依赖：
        pip install -e ".[office]"
    """

    SUPPORTED = (".docx", ".xlsx")

    def parse(self, path: str) -> Document:
        """Parse 函数."""
        p = Path(path)
        suffix = p.suffix.lower()
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix == ".xlsx":
            return self._parse_xlsx(path)
        raise ValueError(f"Unsupported office file: {suffix}")

    def _parse_docx(self, path: str) -> Document:
        try:
            import docx
        except ImportError as exc:
            raise ImportError(
                "DOCX parsing requires 'python-docx'. Install with: pip install -e '.[office]'"
            ) from exc
        document = docx.Document(path)
        title = Path(path).stem
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        tables = []
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            tables.append("\n".join(rows))

        content = "\n\n".join(paragraphs)
        if tables:
            content += "\n\n[Tables]\n" + "\n\n".join(tables)

        file_hash = hashlib.md5(Path(path).read_bytes()).hexdigest()
        return Document(
            doc_id=file_hash,
            title=title,
            source_path=str(Path(path).resolve()),
            file_type="docx",
            total_pages=1,
            chapters=[Chapter(title="全文", start_page=1, end_page=1, level=1)],
            file_hash=file_hash,
            status="pending",
        )

    def _parse_xlsx(self, path: str) -> Document:
        try:
            import openpyxl
        except ImportError as exc:
            raise ImportError(
                "XLSX parsing requires 'openpyxl'. Install with: pip install -e '.[office]'"
            ) from exc
        wb = openpyxl.load_workbook(path, data_only=True)
        title = Path(path).stem
        sheets_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    rows.append(row_text)
            sheets_text.append(f"Sheet: {sheet_name}\n" + "\n".join(rows))

        file_hash = hashlib.md5(Path(path).read_bytes()).hexdigest()
        total = len(wb.sheetnames)
        return Document(
            doc_id=file_hash,
            title=title,
            source_path=str(Path(path).resolve()),
            file_type="xlsx",
            total_pages=total,
            chapters=[
                Chapter(title=sheet, start_page=i + 1, end_page=i + 1, level=1)
                for i, sheet in enumerate(wb.sheetnames)
            ],
            file_hash=file_hash,
            status="pending",
        )
