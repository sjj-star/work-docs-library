"""test_office_parser 模块."""

import docx
import openpyxl
import pytest
from parsers.office_parser import OfficeParser


def test_parse_docx(tmp_path):
    """Test parse docx."""
    doc_path = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("Second paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    doc.save(str(doc_path))

    parser = OfficeParser()
    result = parser.parse(str(doc_path))
    assert result.file_type == "docx"
    assert result.total_pages == 1
    assert len(result.chapters) == 1
    assert result.chapters[0].title == "全文"


def test_parse_xlsx(tmp_path):
    """Test parse xlsx."""
    xlsx_path = tmp_path / "sample.xlsx"
    wb = openpyxl.Workbook()
    ws1 = wb.active
    assert ws1 is not None
    ws1.title = "Sheet1"
    ws1.append(["Name", "Age"])
    ws1.append(["Alice", "30"])
    ws2 = wb.create_sheet("Sheet2")
    ws2.append(["City", "Country"])
    ws2.append(["Beijing", "China"])
    wb.save(str(xlsx_path))

    parser = OfficeParser()
    result = parser.parse(str(xlsx_path))
    assert result.file_type == "xlsx"
    assert result.total_pages == 2
    assert len(result.chapters) == 2
    assert result.chapters[0].title == "Sheet1"
    assert result.chapters[1].title == "Sheet2"


def test_unsupported_suffix(tmp_path):
    """Test unsupported suffix."""
    bad = tmp_path / "bad.pptx"
    bad.write_text("x")
    parser = OfficeParser()
    with pytest.raises(ValueError):
        parser.parse(str(bad))
